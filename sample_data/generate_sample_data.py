"""
Generates a sample RFQ (Excel) and three sample supplier quotations
(Excel, PDF, Word) so you can try the whole app immediately without
needing your own documents. Run:  python sample_data/generate_sample_data.py
"""
from pathlib import Path
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

OUT_DIR = Path(__file__).resolve().parent

# ---------------- RFQ (Excel) ----------------
rfq_header = pd.DataFrame({
    "Field": ["RFQ Ref", "Deadline", "Currency", "Delivery Terms Required", "Payment Terms Required"],
    "Value": ["RFQ-2026-014", "15/09/2026", "INR", "FOB Delhi, within 20 days", "30% advance, balance on delivery"],
})
rfq_items = pd.DataFrame({
    "Item": ["Corrugated Packaging Box (L)", "Steel Rod 12mm", "Industrial Sealant Tube", "Pallet Wrap Film"],
    "Required Quantity": [5000, 2000, 800, 300],
    "Unit": ["pcs", "kg", "pcs", "rolls"],
    "Portion Spec": ["Single unit box", "Per kg bundle", "500ml tube", "20kg roll"],
    "Packaging Spec": ["Bundle of 50", "Bundle of 100kg", "Box of 24", "Shrink pack of 5"],
})
with pd.ExcelWriter(OUT_DIR / "RFQ_2026-014.xlsx", engine="openpyxl") as writer:
    rfq_header.to_excel(writer, sheet_name="RFQ Info", index=False)
    rfq_items.to_excel(writer, sheet_name="Required Items", index=False)

# ---------------- Supplier A quotation (Excel) ----------------
supA_header = pd.DataFrame({
    "Field": ["Supplier", "Quotation Ref", "RFQ Ref", "Delivery Terms", "Payment Terms", "Warranty", "Lead Time (days)"],
    "Value": ["Bharat Packaging Solutions Pvt Ltd", "BPS-Q-991", "RFQ-2026-014",
              "FOB Delhi, within 18 days", "30% advance, balance on delivery", "6 months", "18"],
})
supA_items = pd.DataFrame({
    "Item Name": ["Corrugated Packaging Box (L)", "Steel Rod 12mm", "Industrial Sealant Tube", "Pallet Wrap Film"],
    "Quantity": [5000, 2000, 800, 300],
    "Unit": ["pcs", "kg", "pcs", "rolls"],
    "Unit Price": [18.5, 62.0, 145.0, 410.0],
    "Total Price": [92500, 124000, 116000, 123000],
    "Portion": ["Single unit box", "Per kg bundle", "500ml tube", "20kg roll"],
    "Packaging": ["Bundle of 50", "Bundle of 100kg", "Box of 24", "Shrink pack of 5"],
})
with pd.ExcelWriter(OUT_DIR / "Quotation_BharatPackaging.xlsx", engine="openpyxl") as writer:
    supA_header.to_excel(writer, sheet_name="Quotation Info", index=False)
    supA_items.to_excel(writer, sheet_name="Priced Items", index=False)

# ---------------- Supplier B quotation (Word) ----------------
doc = Document()
doc.add_heading("Supplier Quotation", level=1)
doc.add_paragraph("Supplier: Northline Industrial Supplies")
doc.add_paragraph("Quotation Ref: NIS-2026-77")
doc.add_paragraph("RFQ Ref: RFQ-2026-014")
doc.add_paragraph("Delivery Terms: Ex-works Gurugram, within 25 days")
doc.add_paragraph("Payment Terms: 50% advance, balance on delivery")
doc.add_paragraph("Warranty: 3 months")
doc.add_paragraph("Lead Time: 25 days")
doc.add_paragraph("")
doc.add_paragraph("Note: Sealant tube pricing not included this cycle — to follow separately.")

table = doc.add_table(rows=1, cols=6)
hdr = table.rows[0].cells
for i, h in enumerate(["Item Name", "Quantity", "Unit", "Unit Price", "Portion", "Packaging"]):
    hdr[i].text = h
rows_data = [
    ["Corrugated Packaging Box (L)", "5000", "pcs", "17.90", "Single unit box", "Bundle of 50"],
    ["Steel Rod 12mm", "2000", "kg", "59.50", "Per kg bundle", "Bundle of 100kg"],
    ["Pallet Wrap Film", "300", "rolls", "395.00", "20kg roll", "Shrink pack of 5"],
]
for r in rows_data:
    cells = table.add_row().cells
    for i, val in enumerate(r):
        cells[i].text = val
doc.save(OUT_DIR / "Quotation_Northline.docx")

# ---------------- Supplier C quotation (PDF) ----------------
pdf_path = OUT_DIR / "Quotation_SwiftTrade.pdf"
c = canvas.Canvas(str(pdf_path), pagesize=A4)
width, height = A4
text = c.beginText(40, height - 50)
text.setFont("Helvetica", 11)
lines = [
    "Supplier Quotation",
    "",
    "Supplier: Swift Trade Corp",
    "Quotation Ref: STC-556",
    "RFQ Ref: RFQ-2026-014",
    "Delivery Terms: CIF Delhi, within 15 days",
    "Payment Terms: 20% advance, balance on delivery",
    "Warranty: 12 months",
    "Lead Time: 15 days",
    "",
    "Priced Items:",
    "Corrugated Packaging Box (L)  Qty 5000 pcs  Unit Price Rs. 19.20  Portion: Single unit box  Packaging: Bundle of 50",
    "Steel Rod 12mm  Qty 2000 kg  Unit Price Rs. 61.00  Portion: Per kg bundle  Packaging: Bundle of 100kg",
    "Industrial Sealant Tube  Qty 800 pcs  Unit Price Rs. 139.00  Portion: 500ml tube  Packaging: Box of 24",
    "Pallet Wrap Film  Qty 300 rolls  Unit Price Rs. 405.00  Portion: 20kg roll  Packaging: Shrink pack of 5",
]
for line in lines:
    text.textLine(line)
c.drawText(text)
c.showPage()
c.save()

print("Sample data generated in:", OUT_DIR)
print("- RFQ_2026-014.xlsx")
print("- Quotation_BharatPackaging.xlsx")
print("- Quotation_Northline.docx")
print("- Quotation_SwiftTrade.pdf")
