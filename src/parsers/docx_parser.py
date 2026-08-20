"""
Parses Word (.docx) RFQs / supplier quotations: paragraphs + tables.
"""
from pathlib import Path
from typing import Dict, List
import docx


def parse_docx(file_path: str) -> Dict:
    path = Path(file_path)
    document = docx.Document(str(path))
    text_chunks = [f"Word file: {path.name}"]

    for para in document.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    tables: List[List[List[str]]] = []
    for t_idx, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
        table_text = "\n".join(" | ".join(r) for r in rows)
        text_chunks.append(f"\n[Table {t_idx + 1}]\n{table_text}")

    return {
        "raw_text": "\n".join(text_chunks),
        "tables": tables,
        "source": str(path),
    }
